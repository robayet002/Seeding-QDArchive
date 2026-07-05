"""
Generate the classification report as a Word document (.docx).

Reporting unit ("distribution"): one distribution per
(repository x project type) cell, for the project types QDA_PROJECT and
QD_PROJECT, numbered row-major:

                          Project type
                   QDA_Project     QD_Project
    Repository 1   Distribution 1  Distribution 2
    Repository 2   Distribution 3  Distribution 4
    ...            ...             ...

Document structure:
    - Title + summary
    - "Total Number of Distributions to Report About" overview matrix
    - Per distribution:
        * histogram of primary classes (full class names as bin names,
          count printed on top of each bar)
        * rank-ordered table of the top 20 classes with counts
        * comments on the findings

Run:  python report_docx.py
"""

import io
from collections import Counter

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config import REPORT_DIR
from database import get_connection

REPORT_DOCX_PATH = REPORT_DIR / "classification_report.docx"

PROJECT_TYPES_REPORTED = ["QDA_PROJECT", "QD_PROJECT"]
TOP_N = 20

DARK = "3B3B3B"     # header dark grey (like the mock-up)
MID = "808080"      # mid grey
LIGHT = "D9D9D9"    # light grey cells
BLUE = RGBColor(0x44, 0x72, 0xC4)


# ---------------------------------------------------------------------------
# Data access
# ---------------------------------------------------------------------------

def load_report_data():
    """
    Returns:
        repos:  sorted list of repository ids
        urls:   {repo_id: repository_url}
        cells:  {(repo_id, project_type): {"classes": Counter,
                                           "n_projects": int}}
    """
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("""
        SELECT p.repository_id, p.repository_url, p.type,
               pc.primary_class_name
        FROM projects p
        LEFT JOIN project_classification pc ON pc.project_id = p.id
        WHERE p.type IN ('QDA_PROJECT', 'QD_PROJECT')
    """)

    urls: dict[int, str] = {}
    cells: dict[tuple, dict] = {}

    for repo_id, repo_url, ptype, primary_class in cursor.fetchall():
        urls[repo_id] = repo_url
        cell = cells.setdefault((repo_id, ptype),
                                {"classes": Counter(), "n_projects": 0})
        cell["n_projects"] += 1
        if primary_class:
            cell["classes"][primary_class] += 1

    conn.close()
    repos = sorted(urls)
    return repos, urls, cells


def build_distribution_index(repos):
    """Row-major numbering: repo 1 -> Distribution 1 (QDA), 2 (QD); ..."""
    index = {}
    n = 0
    for repo_id in repos:
        for ptype in PROJECT_TYPES_REPORTED:
            n += 1
            index[(repo_id, ptype)] = n
    return index


# ---------------------------------------------------------------------------
# docx helpers
# ---------------------------------------------------------------------------

def _shade(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _style_cell(cell, text, *, fill=None, bold=False, white=False,
                size=10, center=True):
    cell.text = ""
    para = cell.paragraphs[0]
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    if fill:
        _shade(cell, fill)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    return h


# ---------------------------------------------------------------------------
# Overview matrix ("Total Number of Distributions to Report About")
# ---------------------------------------------------------------------------

def add_distribution_matrix(doc, repos, urls, dist_index):
    _heading(doc, "Total Number of Distributions to Report About", level=1)

    n_rows = 2 + len(repos)          # 2 header rows + one per repository
    table = doc.add_table(rows=n_rows, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row 1: empty | empty | "Project type" (merged over 2 cols)
    top = table.rows[0].cells
    merged = top[2].merge(top[3])
    _style_cell(merged, "Project type", fill=DARK, bold=True, white=True)
    _style_cell(top[0], "", fill="FFFFFF")
    _style_cell(top[1], "", fill="FFFFFF")

    # Header row 2: empty | empty | QDA_Project | QD_Project
    second = table.rows[1].cells
    _style_cell(second[0], "", fill="FFFFFF")
    _style_cell(second[1], "", fill="FFFFFF")
    _style_cell(second[2], "QDA_Project", fill=MID, bold=True, white=True)
    _style_cell(second[3], "QD_Project", fill=MID, bold=True, white=True)

    # Body rows: Repository (merged label) | repo id | Distribution N | N+1
    for i, repo_id in enumerate(repos):
        row = table.rows[2 + i].cells
        _style_cell(row[1], str(repo_id), fill=MID, bold=True, white=True)
        for j, ptype in enumerate(PROJECT_TYPES_REPORTED):
            n = dist_index[(repo_id, ptype)]
            _style_cell(row[2 + j], f"Distribution {n}", fill=LIGHT, bold=True)

    # Merge the "Repository" label column vertically
    if repos:
        label_cell = table.rows[2].cells[0]
        for i in range(1, len(repos)):
            label_cell = label_cell.merge(table.rows[2 + i].cells[0])
        _style_cell(label_cell, "Repository", fill=DARK, bold=True, white=True)

    doc.add_paragraph()
    total = len(repos) * len(PROJECT_TYPES_REPORTED)
    para = doc.add_paragraph()
    run = para.add_run(
        f"Total number of distributions to report about: {total} "
        f"({len(repos)} repository/repositories \u00d7 "
        f"{len(PROJECT_TYPES_REPORTED)} project types). "
        f"Repository URLs: "
        + "; ".join(f"{rid} = {urls[rid]}" for rid in repos)
    )
    run.font.size = Pt(10)


# ---------------------------------------------------------------------------
# Histogram (matplotlib -> PNG -> embedded)
# ---------------------------------------------------------------------------

def _wrap(label: str, width: int = 40) -> str:
    words, lines, current = label.split(), [], ""
    for word in words:
        if len(current) + len(word) + 1 > width:
            lines.append(current)
            current = word
        else:
            current = f"{current} {word}".strip()
    if current:
        lines.append(current)
    return "\n".join(lines)


def make_histogram_png(classes: list[tuple[str, int]], title: str) -> bytes:
    labels = [_wrap(name) for name, _ in classes]
    counts = [count for _, count in classes]

    height = max(3.0, 0.55 * len(classes) + 1.5)
    fig, ax = plt.subplots(figsize=(9.5, height), dpi=200)

    y = range(len(classes))
    bars = ax.barh(y, counts, color="#4472C4", edgecolor="#2F528F")
    ax.set_yticks(list(y))
    ax.set_yticklabels(labels, fontsize=7)
    ax.invert_yaxis()
    ax.set_xlabel("Number of projects", fontsize=8)
    ax.set_title(title, fontsize=10, weight="bold")

    for bar, count in zip(bars, counts):
        ax.text(bar.get_width() + max(counts) * 0.01,
                bar.get_y() + bar.get_height() / 2,
                str(count), va="center", fontsize=8, weight="bold")

    ax.set_xlim(0, max(counts) * 1.12)
    ax.tick_params(axis="x", labelsize=7)
    ax.spines[["top", "right"]].set_visible(False)
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# One distribution section
# ---------------------------------------------------------------------------

def add_distribution_section(doc, n, repo_id, repo_url, ptype, cell):
    doc.add_page_break()
    _heading(doc, f"Distribution {n}: Repository {repo_id} \u2013 {ptype}",
             level=1)
    para = doc.add_paragraph()
    run = para.add_run(f"Repository URL: {repo_url}")
    run.font.size = Pt(9)
    run.font.italic = True

    if cell is None or not cell["classes"]:
        n_projects = cell["n_projects"] if cell else 0
        doc.add_paragraph(
            f"No classified projects in this distribution "
            f"({n_projects} project(s) of type {ptype} found, "
            f"none matched the taxonomy keyword lexicon)."
        )
        return

    classes = cell["classes"].most_common(TOP_N)

    # ---- Histogram -----------------------------------------------------
    _heading(doc, "Histogram of primary classes", level=2)
    png = make_histogram_png(
        classes,
        f"Repository {repo_id}, {ptype}: primary classes (ISIC divisions)",
    )
    doc.add_picture(io.BytesIO(png), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---- Rank-ordered table ---------------------------------------------
    _heading(doc, f"Rank-ordered classes (top {len(classes)})", level=2)
    table = doc.add_table(rows=1 + len(classes), cols=3)
    table.style = "Table Grid"

    header = table.rows[0].cells
    for cell_obj, text in zip(header, ["Rank", "Class (full name)", "Count"]):
        _style_cell(cell_obj, text, fill=DARK, bold=True, white=True,
                    center=(text != "Class (full name)"))

    for rank, (name, count) in enumerate(classes, start=1):
        row = table.rows[rank].cells
        _style_cell(row[0], str(rank), size=9)
        _style_cell(row[1], name, size=9, center=False)
        _style_cell(row[2], str(count), size=9)
        if rank % 2 == 0:
            for c in row:
                _shade(c, "EDF1F8")

    # ---- Comments --------------------------------------------------------
    _heading(doc, "Comments on the findings", level=2)

    n_projects = cell["n_projects"]
    classified = sum(cell["classes"].values())
    dominant, dom_count = cell["classes"].most_common(1)[0]
    n_distinct = len(cell["classes"])
    top3 = cell["classes"].most_common(3)
    top3_share = 100.0 * sum(c for _, c in top3) / classified

    comments = [
        f"This distribution covers {n_projects} project(s) of type {ptype} "
        f"harvested from repository {repo_id}; {classified} of them received "
        f"an ISIC primary class ({100.0 * classified / n_projects:.1f} %).",
        f"The dominant class is {dominant} with {dom_count} project(s).",
        f"{n_distinct} distinct primary classes were identified; the top "
        f"three classes cover {top3_share:.1f} % of all classified projects, "
        f"i.e. the topical distribution is "
        f"{'strongly concentrated' if top3_share > 60 else 'fairly spread out'}.",
        "The classifier is keyword-based over metadata (title, description, "
        "keywords) and base data (file names and readable file content); "
        "multi-word taxonomy keywords are weighted higher as they are more "
        "specific.",
        "Projects without any keyword match remain unclassified; their "
        "metadata was too sparse or too generic to derive an ISIC division "
        "reliably.",
    ]
    for comment in comments:
        doc.add_paragraph(comment, style="List Bullet")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def generate_report(output_path=REPORT_DOCX_PATH):
    repos, urls, cells = load_report_data()

    if not repos:
        print("No QDA_PROJECT/QD_PROJECT data in database - nothing to report.")
        return None

    dist_index = build_distribution_index(repos)

    doc = Document()

    # Base font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("QDArchive \u2013 Classification Report", level=0)
    for run in title.runs:
        run.font.color.rgb = BLUE
    para = doc.add_paragraph(
        "Project types and ISIC topic classification, reported per "
        "distribution (repository \u00d7 project type)."
    )
    para.runs[0].font.italic = True
    doc.add_paragraph()

    add_distribution_matrix(doc, repos, urls, dist_index)

    for repo_id in repos:
        for ptype in PROJECT_TYPES_REPORTED:
            n = dist_index[(repo_id, ptype)]
            cell = cells.get((repo_id, ptype))
            add_distribution_section(doc, n, repo_id, urls[repo_id],
                                     ptype, cell)

    doc.save(output_path)
    print(f"Word report written to: {output_path}")
    return output_path


if __name__ == "__main__":
    generate_report()